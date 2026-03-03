"""
CV Parser: extracts raw text from uploaded PDF or DOCX files.
Validates file type and size before processing. No file is written to disk.
"""

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB


def validate_file(file_bytes: bytes, filename: str) -> tuple[bool, str]:
    """
    Validates upload before parsing. Returns (is_valid, error_message).
    Checks extension and size only — no execution of content.
    """
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"File type '{ext}' is not supported. Upload a PDF or DOCX."
    if len(file_bytes) > MAX_SIZE_BYTES:
        return False, "File exceeds the 5 MB size limit."
    if len(file_bytes) == 0:
        return False, "Uploaded file is empty."
    return True, ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages).strip()
    except ImportError:
        logger.error("PyMuPDF not installed. Run: pip install pymupdf")
        raise RuntimeError("PDF parsing library not available.")
    except Exception as exc:
        logger.error("PDF parsing failed: %s", exc)
        raise RuntimeError("Could not read the PDF file. It may be corrupted or encrypted.") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except ImportError:
        logger.error("python-docx not installed.")
        raise RuntimeError("DOCX parsing library not available.")
    except Exception as exc:
        logger.error("DOCX parsing failed: %s", exc)
        raise RuntimeError("Could not read the DOCX file. It may be corrupted.") from exc


def parse_cv(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Public entry point. Validates then extracts text.
    Returns extracted text string or raises RuntimeError.
    """
    valid, error = validate_file(file_bytes, filename)
    if not valid:
        raise ValueError(error)

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported extension: {ext}")

    if not text or len(text) < 50:
        raise RuntimeError("Extracted CV text is too short. The file may be image-based or empty.")

    return text
